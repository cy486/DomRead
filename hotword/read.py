
# -*- coding: UTF-8 -*-
import docx
import pymysql
#创建数据库
def create():
    # 打开数据库连接
    db = pymysql.connect("localhost", "root", "123456", "hot")

    # 使用 cursor() 方法创建一个游标对象 cursor
    cursor = db.cursor()
    cursor.execute("DROP TABLE IF EXISTS HotWord")
    # 使用预处理语句创建表
    sql = "CREATE TABLE HotWord (id  int primary key auto_increment,word  varchar(100),content longtext, segment longtext,key1 varchar(100),keyword longtext,maincontent longtext,isTec int,isSoup int,isMR int,isMath int,isNews int,baike longtext)"
    cursor.execute(sql)

    # 关闭数据库连接
    db.close()
#插入数据
def insert(one,two):
    # 打开数据库连接
    db = pymysql.connect("localhost", "root", "123456", "hot")

    # 使用cursor()方法获取操作游标
    cursor = db.cursor()
    # SQL 插入语句

    sql_insert = """insert into HotWord(word,content) values('%s','%s')"""%(one,two)
    cursor.execute(sql_insert)

    # 提交到数据库执行
    db.commit()
    # 关闭数据库连接
    db.close()
#主函数把word的内容读出来
def main():
    file=docx.Document("F:\\python-project\\hebei.docx")
    file_word = docx.Document()
    list=[""]
    conlist=[]
    content=""
    #输出段落编号及段落内容
    for p in file.paragraphs:
      if p.style.name == "级别3：黑体 13磅 20行距 段落前后20 左对齐":
       list.append(p.text)
       conlist.append(content)
       content=""
      if p.style.name == "级别4：正文":
         content=content+p.text
    conlist.append(content)
    for i in range(len(list)-1):
        insert(list[i+1],conlist[i+1])

if __name__ == '__main__':
    create()
    main()
