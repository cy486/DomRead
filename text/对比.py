import re

import jieba
import pymysql
from docx import Document
from docx.shared import RGBColor


def cutword(word, baike):
    word = re.sub(r"[0-9\s+\.\!\/_,$%^*()?;；:-【】+\"\']+|[——！，;:。？、~@#￥%……&*（）]+", "", word)
    baike = re.sub(r"[0-9\s+\.\!\/_,$%^*()?;；:-【】+\"\']+|[——！，;:。？、~@#￥%……&*（）]+", "", baike)
    seg_list = jieba.cut(word)
    line = []
    i = 0
    for str in seg_list:
        line.append(str)
    baike_list = jieba.cut(baike)
    line2 = []
    for str in baike_list:
        line2.append(str)
    both = set(line).intersection(set(line2))
    return both


def get_content():
    document = Document()

    document.add_heading('热词查重文档', 0)
    p = document.add_paragraph('查重人员：高得健')
    p.add_run('2019.04.08').italic = True
    p = document.add_paragraph('查重依据：百度文库，维基百科，期刊杂志')
    document.add_paragraph('查重词条数：2500')
    document.add_paragraph('查重率：0.06----%6.2')
    document.add_paragraph('日期二零一九年四月八日')

    conn = pymysql.connect(host="127.0.0.1", user="root", passwd="123456", db="hot", charset="utf8")
    cursor = conn.cursor()
    sql = "select word,content,segment,baike from hotword"
    cursor.execute(sql)
    for row in cursor.fetchall():
        document.add_paragraph(
            row[0], style='List Bullet'
        )
        str=row[2]
        pt=str.split()
        word = cutword(row[1], row[3])
        document.add_paragraph("文库内容：")
        document.add_paragraph(row[3])
        document.add_paragraph("文档内容：")
        p = document.add_paragraph("")
        for oneword in pt:
            if oneword in word:
                run = p.add_run(oneword)
                run.font.color.rgb = RGBColor(250, 0, 0)
            else:
                p.add_run(oneword)
    conn.close()
    document.save('demo2.docx')

if __name__ == '__main__':
    get_content()
